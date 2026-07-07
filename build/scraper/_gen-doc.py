#!/usr/bin/env python3
"""Generate CheapALot Scraper Guide .docx"""

import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('CheapALot — 全网库存爬虫使用指南', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Phase 3: AI 自动抓取外部 B2B 库存网站 → cheapalot.com 自动上架')
run.italic = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

# ── 1. 概述 ──
doc.add_heading('1. 系统概述', 1)
doc.add_paragraph(
    'CheapALot 全网库存爬虫是一套基于 Node.js 的自动化系统, 它能够自动从外部 B2B 库存批发网站 '
    '抓取产品信息、产品图片、价格等数据, 并自动整合到 cheapalot.com 网站, 最终通过 Vercel 自动部署上线.'
)
doc.add_paragraph(
    '整个流程从供应商网站原始 HTML 解析开始, 到 cheapalot.com 三语页面自动生成, 实现了'
    '"零人工操作"的全自动化产品上架链路.'
)

# 核心组件
doc.add_heading('核心组件', 2)
components = [
    ('scraper-core.js', '爬取核心引擎, 支持静态 (Cheerio) 和动态 (Playwright) 两种模式'),
    ('sites.js', '目标站点配置文件, 定义 URL/选择器/分页/分类映射'),
    ('download-images.js', '远程图片下载器, 存到本地 images/products/scraped/'),
    ('normalize.js', '数据归一化器, 含 GPT-4 翻译和简单翻译两种模式'),
    ('run-scraper.js', '主控脚本, 串联所有步骤并触发构建/部署'),
]
for name, desc in components:
    p = doc.add_paragraph()
    run = p.add_run(f'• {name}: ')
    run.bold = True
    p.add_run(desc)

# ── 2. 数据源 ──
doc.add_heading('2. 已配置的数据源', 1)
doc.add_paragraph('系统已预置 5 个 B2B 库存批发网站, 按稳定性分类:')

# Table
table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'
table.cell(0, 0).text = '站点'
table.cell(0, 1).text = '类型'
table.cell(0, 2).text = '状态'
table.cell(0, 3).text = '说明'

sites_data = [
    ('wholesaleclearance.co.uk', '静态', '✅ 已实测', '20 产品/页, 高质量, 已通过完整测试'),
    ('thewholesaler.co.uk', '静态', '🟡 待实测', '结构清晰, 需要微调选择器'),
    ('enviro-stock.co.uk', '动态', '🟡 需 Playwright', 'JS 渲染, 自定义 Gravity Forms 主题'),
    ('palletclearance.app', '动态', '🟡 需 Playwright', 'SPA, 需要无头浏览器'),
    ('eBay / Amazon Liquidation', 'API', '🔴 可扩展', '通过 B2B API 对接, 需独立账号'),
]

for i, (name, typ, status, desc) in enumerate(sites_data, 1):
    table.cell(i, 0).text = name
    table.cell(i, 1).text = typ
    table.cell(i, 2).text = status
    table.cell(i, 3).text = desc

# ── 3. 安装与运行 ──
doc.add_heading('3. 安装与运行', 1)

# 3.1 依赖
doc.add_heading('3.1 依赖安装', 2)
doc.add_paragraph('爬虫依赖以下 npm 包, 已在工作区安装:')
p = doc.add_paragraph()
p.add_run('npm install playwright cheerio').font.name = 'Consolas'
doc.add_paragraph('Playwright 还需要 Chromium 浏览器:')
p = doc.add_paragraph()
p.add_run('npx playwright install chromium').font.name = 'Consolas'

# 3.2 基础运行
doc.add_heading('3.2 基础运行', 2)
doc.add_paragraph('爬取所有已配置站点:')
p = doc.add_paragraph()
p.add_run('node build/scraper/run-scraper.js').font.name = 'Consolas'

doc.add_paragraph('只爬取单个站点:')
p = doc.add_paragraph()
p.add_run('node build/scraper/run-scraper.js wholesaleclearance').font.name = 'Consolas'

doc.add_paragraph('限制每个站点最多 1 页 (测试用):')
p = doc.add_paragraph()
p.add_run('node build/scraper/run-scraper.js wholesaleclearance --max-pages=1').font.name = 'Consolas'

doc.add_paragraph('不执行 git push (本地测试用):')
p = doc.add_paragraph()
p.add_run('node build/scraper/run-scraper.js --no-push').font.name = 'Consolas'

# 3.3 环境变量
doc.add_heading('3.3 环境变量', 2)
env_table = doc.add_table(rows=4, cols=3)
env_table.style = 'Light Grid Accent 1'
env_table.cell(0, 0).text = '变量'
env_table.cell(0, 1).text = '说明'
env_table.cell(0, 2).text = '必需'

envs = [
    ('OPENAI_API_KEY', 'GPT-4 API 密钥, 用于 EN/ES/AR 自动翻译', '可选 (无则用基础翻译)'),
    ('GIT_AUTO_PUSH', '设为 "1" 自动 git push 到 GitHub', '可选 (默认不推送)'),
    ('NODE_PATH', 'node_modules 路径, 因依赖安装在隔离目录', '必需 (运行脚本时设置)'),
]

for i, (var, desc, req) in enumerate(envs, 1):
    env_table.cell(i, 0).text = var
    env_table.cell(i, 1).text = desc
    env_table.cell(i, 2).text = req

# ── 4. 完整流程 ──
doc.add_heading('4. 完整运行流程', 1)

steps = [
    ('Step 1: 发起 HTTP 请求', '使用 Node.js https/http 模块拉取目标站点 HTML (静态) 或使用 Playwright 无头浏览器 (动态).'),
    ('Step 2: DOM 解析', '使用 cheerio 加载 HTML, 按 sites.js 中定义的选择器提取产品数据.'),
    ('Step 3: 字段提取', '产品名/价格/Job Lot 价格/折扣%/SKU/图片 URL 全部解析成结构化 JSON.'),
    ('Step 4: 自动分类', '根据产品标题中的关键词自动归类到 8 个类别: household, electrical, apparel, diy, furniture, toys, sports, mixed.'),
    ('Step 5: 图片下载', '把所有产品图片 (包括懒加载 data-src) 下载到本地 images/products/scraped/ 目录.'),
    ('Step 6: GPT-4 翻译', '调用 GPT-4o-mini API, 把产品标题翻译成西班牙语和阿拉伯语 (如 API 不可用, 自动降级为关键词翻译).'),
    ('Step 7: 合并数据', '把爬取的新产品合并到 data/products.json, 自动分配 p13/p14... 序列号, 保留手动添加的产品.'),
    ('Step 8: 构建 HTML', '调用 build-products.js 重新生成 EN/ES/AR 三语版本的 products.html, 含筛选侧边栏和 deals 区域.'),
    ('Step 9: Git push 部署', 'git commit + git push origin main → Vercel 自动部署, 1-2 分钟内 cheapalot.com 上线.'),
]

for i, (title_str, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.add_run(f'{title_str}').bold = True
    p.add_run(f' — {desc}')

# ── 5. 实测结果 ──
doc.add_heading('5. 实测运行结果', 1)
doc.add_paragraph('针对 wholesaleclearance.co.uk 单站点单页实测结果:')

stats = doc.add_table(rows=8, cols=2)
stats.style = 'Light Grid Accent 1'
stats.cell(0, 0).text = '指标'
stats.cell(0, 1).text = '数值'

results = [
    ('爬取产品数', '20 个 job lot'),
    ('产品图片下载', '20/20 全部成功 (100%)'),
    ('产品分类分布', 'apparel: 10, mixed: 4, toys: 2, electrical: 1, 其它: 3'),
    ('价格区间', '£0.32 — £12.31 (单件), Job Lot 价格最高 £4,420'),
    ('数据合并', 'products.json 从 12 增至 32 个产品'),
    ('HTML 重新生成', 'products.html (EN/ES/AR) 全部成功, 平均 1000+ 行'),
    ('翻译模式', 'GPT-4 不可达, 自动降级到简单翻译 (基础关键词映射)'),
]

for i, (k, v) in enumerate(results, 1):
    stats.cell(i, 0).text = k
    stats.cell(i, 1).text = v

# ── 6. 添加新站点 ──
doc.add_heading('6. 添加新站点', 1)
doc.add_paragraph('在 build/scraper/sites.js 中添加新站点配置, 步骤如下:')

doc.add_paragraph('Step 1: 浏览器打开目标网站, 用开发者工具定位产品卡片的 CSS class (如 .product, .joblot 等).')
doc.add_paragraph('Step 2: 在 sites.js 的 module.exports 中添加新键名, 配置:')

p = doc.add_paragraph()
run = p.add_run('''newsite: {
  name: '站点名',
  baseUrl: 'https://...',
  startUrl: 'https://.../products',
  type: 'static',  // 或 'dynamic'
  maxPages: 3,
  delayMs: 1500,
  paginationParam: 'page',  // 或 'paginationPath' for WordPress
  selectors: {
    productCard: '.product-card',
    title: '.product-title',
    price: '.price',
    image: 'img',
    imageUrlAttr: 'data-src',
    link: 'a',
    linkAttr: 'href',
  },
  categoryMap: {
    'fashion|clothing': 'apparel',
    'electronic|phone': 'electrical',
    // ...
  },
},''')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph('Step 3: 测试爬取, 验证选择器正确:')
p = doc.add_paragraph()
p.add_run('node build/scraper/run-scraper.js newsite --max-pages=1 --no-push').font.name = 'Consolas'

# ── 7. 常见问题 ──
doc.add_heading('7. 常见问题与解决', 1)

faq = [
    ('Q: 爬取到 0 个产品怎么办?',
     'A: 检查 selectors 是否匹配目标站点的 DOM 结构. 使用浏览器开发者工具查找产品卡片的真实 CSS class, 更新 sites.js.'),
    ('Q: 图片下载失败?',
     'A: 1) 检查 URL 是否有效; 2) 部分网站有防盗链, 需要在 download-images.js 中添加 Referer 头; 3) 跳过 placeholder 图片.'),
    ('Q: GPT-4 翻译超时?',
     'A: 10 秒超时后自动降级到简单翻译 (基础关键词映射). 翻译质量会下降, 但不影响系统运行. 可设置 OPENAI_API_KEY 启用完整翻译.'),
    ('Q: Cloudflare 拦截爬取?',
     'A: 如 merkandi.co.uk, 需要更高级的反爬策略: 1) 启用 Playwright (type: dynamic); 2) 随机 User-Agent; 3) 添加 Cookie 模拟真实浏览器.'),
    ('Q: Git push 失败?',
     'A: 默认不自动 push, 用 --no-push 标志. 手动 push 时如遇代理问题: env -u HTTP_PROXY -u HTTPS_PROXY git push origin main.'),
    ('Q: 想加入价格过滤 (只爬 £1 以下的商品)?',
     'A: 在 scraper-core.js 的 parseProducts 循环中添加: if (parsePrice(price) > 1) return; 跳过高于阈值的商品.'),
]

for q, a in faq:
    p = doc.add_paragraph()
    p.add_run(q).bold = True
    p2 = doc.add_paragraph(a)

# ── 8. 成本与扩展建议 ──
doc.add_heading('8. 成本与扩展建议', 1)

doc.add_heading('当前运行成本', 2)
cost_table = doc.add_table(rows=6, cols=2)
cost_table.style = 'Light Grid Accent 1'
cost_table.cell(0, 0).text = '项目'
cost_table.cell(0, 1).text = '成本'

costs = [
    ('静态站点爬取 (Cheerio)', '£0 (免费)'),
    ('动态站点爬取 (Playwright)', '£0 (免费, 但 Vercel Serverless 不适合)'),
    ('GPT-4 翻译 (gpt-4o-mini)', '£40-80/月 (假设每天 100 个新产品)'),
    ('图片存储 (Vercel)', '免费 (前 100GB 流量)'),
    ('总计', '£0-80/月'),
]
for i, (k, v) in enumerate(costs, 1):
    cost_table.cell(i, 0).text = k
    cost_table.cell(i, 1).text = v

doc.add_heading('扩展建议', 2)
extensions = [
    '1. 加入定时任务: 用 GitHub Actions + cron 每天自动爬取, 完全自动化.',
    '2. 加入去重逻辑: 用 SKU 或产品名 hash 检测已有产品, 避免重复上架.',
    '3. 加入价格监控: 同一产品跨站点比价, 自动选择最低价, 增加竞争力.',
    '4. 加入智能筛选: 利润分析 (售价 - 成本), 自动过滤低利润产品.',
    '5. 对接 eBay/Amazon Liquidation API: 增加更多数据源.',
    '6. 加入图片优化: 用 sharp.js 压缩图片到 WebP, 减小页面体积.',
    '7. 增量爬取: 只爬取新增产品, 用 lastProductId 记录上次位置.',
]
for ext in extensions:
    p = doc.add_paragraph()
    p.add_run('• ').bold = True
    p.add_run(ext)

# ── 9. 相关文件清单 ──
doc.add_heading('9. 相关文件清单', 1)

files = [
    'build/scraper/sites.js — 站点配置',
    'build/scraper/scraper-core.js — 爬取核心',
    'build/scraper/download-images.js — 图片下载',
    'build/scraper/normalize.js — 数据归一化 + 翻译',
    'build/scraper/run-scraper.js — 主控脚本',
    'data/products.json — 产品数据 (32 个产品)',
    'images/products/scraped/ — 60 张已下载的产品图',
    'products.html / es/ / ar/ — 三语产品页面',
]

for f in files:
    p = doc.add_paragraph()
    p.add_run('• ' + f).font.name = 'Consolas'

# Save
output_path = sys.argv[1] if len(sys.argv) > 1 else 'CheapALot-Scraper-Guide.docx'
doc.save(output_path)
print(f'Generated: {output_path}')
